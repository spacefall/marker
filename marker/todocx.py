from marker.parser import fromfile
from marker.classes import *
from docx import Document  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore


def todocx(parsed_mkr: list[MarkerTypes], output_filename: str) -> None:
    document: Document = Document()
    paragraph: Paragraph | None = None

    # this is a very manual way of doing this, a more dynamic way would be better
    # but that would require a failsafe for when the script doesn't recognize the tag
    # so I think it's better to leave it like this for now
    infodict: dict[str, bool] = {
        k: False
        for k in [
            "h",  # header
            "b",  # bold
            "i",  # italic
            "s",  # strikethrough
            "u",  # underline
        ]
    }

    for elem in parsed_mkr:
        match elem:
            case Header():
                # adds heading
                paragraph = document.add_heading(level=elem.level)
                infodict["h"] = True
            case Bold():
                infodict["b"] = not infodict["b"]
            case Italic():
                infodict["i"] = not infodict["i"]
            case Strikethrough():
                infodict["s"] = not infodict["s"]
            case Underline():
                infodict["u"] = not infodict["u"]
            case Newline():
                if paragraph is not None:
                    # resetting paragraph
                    paragraph = None
                else:
                    # adding new paragraph (that results only in a new line since no text is specified) if no paragraph
                    document.add_paragraph()

                # resetting every bool in dict
                for k in infodict:
                    infodict[k] = False
            case str():
                # new paragraph if needed
                if paragraph is None:
                    paragraph = document.add_paragraph()
                # add to current paragraph
                run = paragraph.add_run(elem)

                # bold
                if infodict["b"]:
                    run.bold = True
                # italics
                if infodict["i"]:
                    run.italic = True
                # strikethrough
                if infodict["s"]:
                    run.font.strike = True
                # underline
                if infodict["u"]:
                    run.font.underline = True

    if output_filename[-5:] != ".docx":
        output_filename += ".docx"
    document.save(output_filename)


if __name__ == "__main__":
    todocx(fromfile("test.mkr"), "test.docx")
